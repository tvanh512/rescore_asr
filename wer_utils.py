import re
import string
import csv
import Levenshtein
import jiwer
from jiwer.process import WordOutput

import pandas as pd
import numpy as np
from pathlib import Path
import os
import math
from decimal import InvalidOperation
import contractions


######################
# text formatting and WER functions
######################

def strip_punct(instr):
    newstr = ''
    for word in instr.split():
	    # delete punct
        word = word.strip(string.punctuation)

        # delete commas inside numbers
        m = re.match(r'(\d*),(\d)', word)
        if m != None:
            word = word.replace(',', '')

        # commas inside words become space
        word = re.sub(",", " ", word)

        # hyphens inside words become space
        word = re.sub("-", " ", word)
        word = word.strip()

        newstr += ' ' + word
    newstr = newstr.strip()
    return newstr

def remove_in_brackets(text):
    # removes any clause in brackets or parens, and the brackets themselves
    return re.sub("[\(\[\<].*?[\)\]\>]+", " ", text)

def caught_num2words(text):
    from num2words import num2words
    # first do currency replacements
    if '$' in text:
        text = re.sub('\$([0-9]+)', '\g<1> dollars', text)
    # strip punctuation 
    text=strip_punct(text)
    # catch strings that might be converted to infinity or NaN and return as is... 
    naughty_words = ['INF','Inf','inf','NAN','NaN', 'nan', 'NONE','None','none', "Infinity", "INFINITY", "infinity"]
    if text in naughty_words:
        return text
    try:
        if len(text.split()) > 1:
            return ' '.join([caught_num2words(word) for word in text.split()])
        else:
            return num2words(text)
    except (InvalidOperation, ValueError) as error:
        return text
    
def expand_contractions(str):
        expanded_words = []
        for wrd in str.split():
            expanded_words.append(contractions.fix(wrd))
        str = ' '.join(expanded_words)
        return str

def format_text_for_wer(text):
    # function to format text or lists of text (e.g. asr, transcript) for wer computation. 
    # Converts from list to a single string and apply some text normalization operations
    # note that the clean_REV_transcript function should be applied first to remove REV-specific keywords 
    # and extract text from docx format tables
    
    if isinstance(text,list):
        text = ' '.join(text)
    text = text.replace('\n',' ') # replace newline with space
    text = remove_in_brackets(text) # removes non-spoken annotations such as [inaudible]
    text = re.sub('%\w+','', text) # remove %HESITATION etc
    text = ' '.join([caught_num2words(str) for str in text.split(' ')]) # spell out numbers
    text = expand_contractions(text)
    text = strip_punct(text)
    text = text.lower()
    text = re.sub('\s+',' ',text) # replace multiple space with single

    return text

def clean_REV_transcript(docx_fname, txt_fname):
    import docx
    doc = docx.Document(docx_fname)
    doctext = [p.text for p in doc.paragraphs]
    # The transcript may be packed into tables
    for t in doc.tables:
        for c in t._cells:
            doctext.append(c.text)

    # write unstripped transcript to .txt
    txt_fname_diarized = re.sub('.txt','_diarized.txt',txt_fname)
    with open(txt_fname_diarized,'w') as outfile:
        outfile.write('\n'.join(doctext))

    # strip the various Speaker IDs and crosstalk indicators  off
    doc_stripped = [re.sub('Speaker \d+:','',line) for line in doctext]
    doc_stripped = [re.sub('.+:','',line) for line in doc_stripped] # remove anything before colon - this is speaker ID
    doc_stripped = [re.sub(r"\t",'',line) for line in doc_stripped] # remove tabs
    doc_stripped = [line  for line in doc_stripped if not re.match(r'^\s*$', line)] # remove blank lines
    doc_stripped = [remove_in_brackets(line) for line in doc_stripped] # remove sections in brackets or parens
    doc_stripped = [strip_punct(line)  for line in doc_stripped] # remove punct
    # write stripped transcript to txt
    with open(txt_fname,'w') as outfile:
        outfile.write('\n'.join(doc_stripped))

def HHMMSS_to_sec(time_str):
    """Get Seconds from timestamp string with milliseconds."""
    if not time_str:
        return None
    if time_str.count(':')==2:
        h, m, s = time_str.split(':')
    elif time_str.count(':')==3:
    # weird timestamps where there is a field followign seconds delimited by colon
        h, m, s, u = time_str.split(':')
        # determine whether ms field is in tenths or hundredths or thousandths by countng how many digits
        if len(u)==1:
            print('Weird time format detected - HH:MM:SS:tenths - please verify this is how you want the time interpreted')
            ms = float(u)/10
        elif len(u)==2: # hundredths
            ms = float(u)/100
        elif len(u)==3: # hundredths
            ms = float(u)/1000
        else:
            print(f'input string format not supported: {time_str}')
            return None
        s = int(s)+ms
    elif time_str.count(':')==1:
        # print('missing HH from timestamp, assuming MM:SS')
        m, s = time_str.split(':')
        h=0
    else:
        print(f'input string format not supported: {time_str}')
        return None
    return int(h) * 3600 + int(m) * 60 + float(s) 

def name_counter(str):
    patterns = ['\[redacted[\w\s]*\]', '\[Student \d+\]+']
    N = 0
    for p in patterns:
        N+=len(re.findall(p, str,re.IGNORECASE))
    return N

NullWordOutput=WordOutput(references=[[]],hypotheses=[[]],alignments=[[]],wer=None, 
mer=None, 
wil=None, 
wip=None, 
hits=None, 
substitutions=None, 
insertions=None, 
deletions=None)


def align_words(ref,hyp):
    '''
    Aligns two lists of words and outputs the alignment and edit operations
        Parameters:
            ref: reference string
            hyp: hypothesis string


        Returns:
            aligned: a pandas dataframe representing the alignment, 1 row per word 
                with columns:
                    ref_ix: index of word in the reference 
                    hyp_ix: index of word in the hypothesis
                    reference: word from the reference
                    hypothesis: matched word in hypothesis
                    operation: symbolic representations of operation 
                        {'=' : match, 
                        '+':insertion,
                        '-' : deletion,
                        '<>' : substitution
                        }
                    index_edit_ops: index into the edit_ops variable pertaining to each row 
            edit_ops: data frame of word-level operations to go from ref -> hyp

    
    '''

    # get all words and encode as UTF-8 characters to get alignment operations at word-level
    lexicon = list(set(ref+hyp))
    word2digit = dict((lexicon[i],chr(i)) for i in range(0,len(lexicon)))
    asr_uni =  [word2digit[w] for w in hyp]
    trans_uni =  [word2digit[w] for w in ref]
    edit_ops = pd.DataFrame(Levenshtein.editops(''.join(trans_uni),''.join(asr_uni)),
        columns = ['operation','transcript_ix','asr_ix'])
    

    # align the sequences, starting with a dumb alignment where they start together, then inserting as necessary
    aligned_ref = ref.copy()
    aligned_hyp = hyp.copy()
    ix_edit_ops = [np.NaN] *len(aligned_ref)
    aligned_ops =['='] *len(aligned_ref)
    aligned_ref_ix = list(range(len(ref)))
    aligned_hyp_ix = list(range(len(hyp)))

    ins_count = 0 # counter for insertion operations which increase the length of the ref seq thus change the indices
    del_count = 0 # counter for deletion operations which increase the length of the hyp seq thus change the indices
    for [i,ops] in edit_ops.iterrows():
        if ops['operation'] == 'insert':
            aligned_ref.insert(ins_count+ops['transcript_ix'],'_')
            aligned_ops.insert(ins_count+ops['transcript_ix'],'ins')
            aligned_ref_ix.insert(ins_count+ops['transcript_ix'],None)
            ix_edit_ops.insert(ins_count+ops['transcript_ix'],i)
            ins_count = ins_count+1

        if ops['operation'] == 'delete':
            aligned_hyp.insert(del_count+ops['asr_ix'],'_')
            aligned_ops[ins_count + ops['transcript_ix']] = 'del'
            aligned_hyp_ix.insert(del_count+ops['asr_ix'],None)
            ix_edit_ops[ins_count + ops['transcript_ix']] = i
            del_count=del_count+1

        if ops['operation'] == 'replace':
            aligned_ops[ins_count+ ops['transcript_ix']] ='sub' 
            ix_edit_ops[ins_count+ ops['transcript_ix']] =i
           

    aligned = pd.DataFrame({
        'ref_ix':aligned_ref_ix,
        'hyp_ix':aligned_hyp_ix,
        'reference':aligned_ref,
        'hypothesis' : aligned_hyp,
        'operation' : aligned_ops,
        'index_edit_ops' : ix_edit_ops
        })
    # convert columns that shuold be integer to the nullable integer Int64 (uppercase I)
    aligned = aligned.astype({'ref_ix':'Int32','hyp_ix':'Int32','index_edit_ops':'Int32'})
    return aligned, edit_ops


def wer_from_counts(N, sub_count, del_count, ins_count):
    '''
    Computes WER and related measures from edit operation counts and reference wordcount
    Useful to recompute measures without needing to realign raw text
    '''
    if N==0: # WER undefined because no words in reference
        meas = {
                'wer': math.nan,
                'mer': math.nan,
                'hits': math.nan,
                'sub_rate': math.nan,
                'del_rate': math.nan,
                'ins_rate': math.nan
                }
    else:
        meas = {
                'wer': (sub_count + del_count + ins_count)/N,
                'mer': 1 - (N - sub_count - del_count)/N,
                'hits': N - sub_count - del_count ,
                'sub_rate': sub_count/N,
                'del_rate': del_count/N,
                'ins_rate': ins_count/N
                }
    return meas


def wer_from_csv(csv_path, 
        refcol='ref', 
        hypcol='hyp', 
        IDcol='ID',
        files_to_include=None,
        return_alignments=False, 
        normalise = True, 
        printout=True):
    # compute WER with just  isat normalisation, i.e. format_text_for_wer
    res = pd.read_csv(csv_path).astype(str)
    # TODO: implement index by col number
    refs=res[refcol]
    hyps = res[hypcol]

    if files_to_include:
        assert isinstance(files_to_include,list) ,'files_to_include should be a list'
        mask = res[IDcol].isin(files_to_include)
        refs=refs[mask]
        hyps=hyps[mask]
        print(f'Including {len(hyps)} hypotheses matching files_to_include...')

    if normalise:
        refs=refs.apply(format_text_for_wer)
        hyps=hyps.apply(format_text_for_wer)

    #ID,ref,hyp,ref_norm,hyp_norm
    wer_meas = jiwer.compute_measures(list(refs), list(hyps))

    if not return_alignments:
        # remove alignments 
        del wer_meas['ops']
        del wer_meas['truth']
        del wer_meas['hypothesis']
    wer_meas['word_counts'] = wer_meas['substitutions']+wer_meas['deletions']+wer_meas['hits']
    wer_meas['sub_rate'] = wer_meas['substitutions']/wer_meas['word_counts'] 
    wer_meas['del_rate'] = wer_meas['deletions']/wer_meas['word_counts'] 
    wer_meas['ins_rate'] = wer_meas['insertions']/wer_meas['word_counts'] 

    if printout:
        for key in ['wer','sub_rate','del_rate','ins_rate']:
            print((f"{key}={wer_meas[key]:.2f}" ))

        table_string=f"{wer_meas['wer']:.2f} ({wer_meas['sub_rate']:.2f}, {wer_meas['del_rate']:.2f}, {wer_meas['ins_rate']:.2f})"
        print('WER (s,d,i):')
        print(table_string)
    return wer_meas

def wer_dirs(ref_dir, 
            hyp_dir, 
            ref_ext = '.txt', 
            hyp_ext = '.asr.txt', 
            files_to_include=None,
            normalize=True):
    """Run WER compoutation where REF and HYP files have matching filenames and directory structure
    Args:
        ref_dir (str): path to dir containing references
        hyp_dir (str): path to dir containing hypotheses
        ref_ext (str, optional): file extension of reference. Defaults to '.txt'.
        hyp_ext (str, optional): file extension of hypothesis. Defaults to '.asr.txt'.
        normalize (bool, optional): whether to normalize text before WER computation
    Returns:
        metrics (dict): WER metrics
    """    
    # TODO: write a more flexible search for matching hyps that doesnt require matching dir structure

    # ensure ref and hyp dirs in consistent format
    ref_dir = os.path.abspath(os.path.expanduser(ref_dir))
    hyp_dir = os.path.abspath(os.path.expanduser(hyp_dir))

    if files_to_include:
        assert isinstance(files_to_include,list) ,'files_to_include should be a list'

        hyplist=[]
        for f in list(Path(hyp_dir).rglob(f"*{hyp_ext}")):
            ID = str(Path(f).name).replace(hyp_ext,'')
            print(f)
            print(f"ID: {ID}")
            hypf=str(f)
            # hypf = os.path.join(hyp_dir,f)
            if ID in files_to_include:
                # if ''.join(Path(f).suffixes) == hyp_ext:
                hyplist.append(hypf)
        print(f'Including {len(hyplist)} hypotheses matching files_to_include...')

    else:
        hyplist =  list(Path(hyp_dir).rglob(f"*{hyp_ext}"))
    n=[]
    s=[]
    d=[]
    i=[]
    n_utt=0
    for asr_file in hyplist:
        asr_file=str(asr_file) # cos it maybe was a Path 
        # print(ref_file)
        ref_file = asr_file.replace(hyp_dir, ref_dir).replace(hyp_ext, ref_ext)
        asr = open(asr_file,'r').read()

        try:
            ref = open(ref_file,'r').read()
        except FileNotFoundError:
            print(f'reference does not exist: {ref_file}')
        if normalize:
            asr = format_text_for_wer(asr)
            ref = format_text_for_wer(ref)

        if not ref: # empty ref
            print(f'reference exists but is empty: {ref_file}')
            continue
        wer_meas = jiwer.process_words(ref, asr)
        n.append(len(ref.split(' ')))
        s.append(wer_meas.substitutions)
        d.append(wer_meas.deletions)
        i.append(wer_meas.insertions)
        n_utt+=1
    metrics = wer_from_counts(sum(n),sum(s),sum(d),sum(i))
    metrics['n_ref']=n_utt
    return metrics

def csv_to_asrtxt(in_csv, out_dir):
    os.makedirs(out_dir, exist_ok=True)
    df=pd.read_csv(in_csv)
    for i,row in df.iterrows():
        # print(row['ref'])
        # print(row['hyp'])
        # print(row['ID'])
        with open(os.path.join(out_dir, f"{row['ID']}.asr.txt" ),'w') as f:
            f.write(str(row['hyp']))

def asrtxt_to_csv(out_csv,
                ref_dir,
                hyp_dir,
                ref_ext = '.txt',
                hyp_ext = '.asr.txt',
                files_to_include=None,
):
  
    with open(out_csv, 'w', newline='') as csvfile:
        fieldnames = ['ID',
                    'ref',
                    'hyp',
                    'ref_norm',
                    'hyp_norm']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames,delimiter=',')
        writer.writeheader()

        # ensure ref and hyp dirs in consistent format
        ref_dir = os.path.abspath(os.path.expanduser(ref_dir))
        hyp_dir = os.path.abspath(os.path.expanduser(hyp_dir))

        if files_to_include:
            assert isinstance(files_to_include,list) ,'files_to_include should be a list'

            reflist=[]
            for f in os.listdir(ref_dir):
                reff = os.path.join(ref_dir,f)
                if str(Path(f).stem.split('.')[0]) in files_to_include:
                    if ''.join(Path(f).suffixes) == ref_ext:
                        reflist.append(reff)
            print(f'Including {len(reflist)} references matching files_to_include...')
        else:
            reflist =  [os.path.join(ref_dir,f) for f in os.listdir(ref_dir) if ''.join(Path(f).suffixes) == ref_ext]
        
        for ref_file in reflist:
            # print(ref_file)
            ID=Path(ref_file).stem
            asr_file = ref_file.replace(ref_dir, hyp_dir).replace(ref_ext, hyp_ext)
            asr = open(asr_file,'r').read()
            ref = open(ref_file,'r').read()

            asr_norm = format_text_for_wer(asr)
            ref_norm = format_text_for_wer(ref)
            
            writer.writerow( {
            'ID':ID,
            'ref':ref,
            'hyp':asr,
            'ref_norm':ref_norm,
            'hyp_norm':asr_norm})